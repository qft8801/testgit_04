#!/usr/bin/python
#-*-coding:utf-8 -*-
#创建并写入word文档
import docx

#创建内存中的word文档对象
file=docx.Document()

#写入若干段落
file.add_paragraph("啊")
file.add_paragraph("山，真的很高")
file.add_paragraph("高的吓人")
file.add_paragraph("啊！！")
file.add_paragraph("实在是高")

#保存
file.save("D:\\Office_01\\writeResult.docx")