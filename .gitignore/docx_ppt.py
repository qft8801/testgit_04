
#-*-coding:utf-8 -*-

import docx
import re
#获取文档对象
file=docx.Document("D:\Office_01\word题库.docx")
print("段落数:"+str(len(file.paragraphs)))#段落数为13，每个回车隔离一段
print("---------------------------------------------------------------------------")
arr=[]
for i in range(len(file.paragraphs)):
	arr.append(file.paragraphs[i].text)
print(arr[1:36])

